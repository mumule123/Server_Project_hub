from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os,re
from datetime import datetime
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
from config import WORD_REPORT_NAME, PDF_REPORT_NAME # 配置文件

current_dir = os.path.dirname(os.path.abspath(__file__))

# 使用多种方法转换 Word 到 PDF
def convert_word_to_pdf(input_docx, output_pdf):
    """
    尝试使用多种方法将 Word 文档转换为 PDF
    """
    print(f"尝试将 {input_docx} 转换为 {output_pdf}")
    
    # 方法 1: 使用 docx2pdf
    try:
        from docx2pdf import convert
        print("使用 docx2pdf 进行转换...")
        convert(input_docx, output_pdf)
        
        # 验证文件是否成功生成且大小合理
        if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 1000:
            print(f"成功使用 docx2pdf 转换为 PDF: {output_pdf}")
            return True
        else:
            print("PDF 生成但可能无效，尝试其他方法")
    except Exception as e:
        print(f"docx2pdf 转换失败: {e}")
    
    # 方法 2: 使用 win32com (仅 Windows 平台)
    try:
        import sys
        if sys.platform == 'win32':
            print("尝试使用 win32com 进行转换...")
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # 转换路径为绝对路径
            abs_input_path = os.path.abspath(input_docx)
            abs_output_path = os.path.abspath(output_pdf)
            
            doc = word.Documents.Open(abs_input_path)
            doc.SaveAs(abs_output_path, FileFormat=17)  # 17 代表 PDF 格式
            doc.Close()
            word.Quit()
            
            # 验证文件
            if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 1000:
                print(f"成功使用 win32com 转换为 PDF: {output_pdf}")
                return True
            else:
                print("PDF 生成但可能无效，尝试其他方法")
    except Exception as e:
        print(f"win32com 转换失败: {e}")
    
    # 方法 3: 使用库 comtypes (Windows 备选方案)
    try:
        if sys.platform == 'win32':
            print("尝试使用 comtypes 进行转换...")
            import comtypes.client
            
            abs_input_path = os.path.abspath(input_docx)
            abs_output_path = os.path.abspath(output_pdf)
            
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            
            doc = word.Documents.Open(abs_input_path)
            doc.SaveAs(abs_output_path, FileFormat=17)
            doc.Close()
            word.Quit()
            
            if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 1000:
                print(f"成功使用 comtypes 转换为 PDF: {output_pdf}")
                return True
            else:
                print("PDF 生成但可能无效，尝试其他方法")
    except Exception as e:
        print(f"comtypes 转换失败: {e}")
    
    # 如果所有方法都失败，返回 False
    print("所有 PDF 转换方法均失败")
    return False

def create_government_blockchain_report(
    stored_code_values,
    data,
    output_file=os.path.join(current_dir,'document',WORD_REPORT_NAME)
):
    # 若 data 是字符串则先解析
    if isinstance(data, str):
        try:
            data = json.loads(data)

            # 如果解析后是 list[str]，还要进一步把每个 entry 再做 json.loads
            if isinstance(data, list) and all(isinstance(d, str) for d in data):
                data = [json.loads(d) for d in data]

        except Exception as e:
            print(f"解析 data 出错，无法解析为 JSON：{e}")
            return

    # 创建新文档
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


    # 设置正文中文字体为宋体，英文字体为Times New Roman
    def set_run_font(run, size=Pt(12), bold=False):
        font = run.font
        font.size = size
        font.bold = bold

        # 设置英文字体
        run.font.name = "Times New Roman"

        # 设置中文字体为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 辅助函数：添加带有粗体标签的段落
    def add_bold_paragraph(doc, field_title, description):
        p = doc.add_paragraph()
        run1 = p.add_run(f"{field_title}：")
        set_run_font(run1, bold=True)
        run2 = p.add_run(description)
        set_run_font(run2)

    # 辅助函数：添加字段值
    def add_field_value(doc, label, value):
        p = doc.add_paragraph()
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, bold=True)
        value_run = p.add_run(str(value))
        set_run_font(value_run)

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 添加页眉标题
    header_para = doc.sections[0].header.paragraphs[0]
    header_run = header_para.add_run("源代码自研率检测报告")
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header_run, size=Pt(12))

    # 添加页脚和页码
    def add_page_number(document):
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 创建页码文本
        page_num_run = footer_para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        # 设置页码字体
        font = page_num_run.font
        font.name = "Times New Roman"
        font.size = Pt(10)

    # 添加页码
    add_page_number(doc)

    # === 第一页：封面 ===
    # 添加主标题（居中）
    for _ in range(8):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = main_title.add_run("源代码自研率检测报告")
    set_run_font(title_run, size=Pt(28), bold=True)
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑") # 这个标题单独设置成微软雅黑吧

    # 底部信息
    for _ in range(12):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    bottom_info = doc.add_paragraph()
    bottom_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_time = datetime.now().strftime("%Y 年 %m 月 %d 日 %H:%M:%S")

    bottom_run_first = bottom_info.add_run("编制单位：佛山大学\n")
    set_run_font(bottom_run_first, size=Pt(14))
    bottom_run_second = bottom_info.add_run(f"编制时间：{current_time}")
    set_run_font(bottom_run_second, size=Pt(14))


    # === 第二页：内容 ===
    # 设置标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("源代码自研率检测报告")
    set_run_font(title_run, size=Pt(16), bold=True)
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 添加简介文本
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "    本报告对上传代码自研率进行了检测，包括代码同源分析、代码引用、代码行级开源占比和代码文件级开源占比。"
    )
    set_run_font(intro_run)

    # 添加定义标题
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. 定义")
    set_run_font(h1_run, bold=True)

    # 添加项目简介内容
    intro_text = """    源代码安全审计的主要目的是提高源代码质量，通过对程序源代码进行检查和分析，发现源代码在软件设计、测试、应用部署等各阶段中可能存在的安全缺陷或安全漏洞，从源头上避免潜在的安全风险。"""
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(intro_text)
    set_run_font(intro_run)

    # 一、字段说明
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("2. 字段说明")
    set_run_font(h1_run, bold=True)

    para = doc.add_paragraph()
    run = para.add_run("以下是检测系统输出的字段说明：")
    set_run_font(run, size=Pt(12))  # ✅ 显式设置


    add_bold_paragraph(doc, "待检测代码路径", "用户上传、参与查重的代码文件路径。")
    add_bold_paragraph(doc, "参考代码路径", "系统比对的原始或参考源代码路径。")
    add_bold_paragraph(doc, "待检测文件的相似度得分", "该文件中与参考代码相似的比例（0~1）。")
    add_bold_paragraph(doc, "参考文件的相似度得分", "参考代码中出现在待检测文件中的比例（0~1）。")
    add_bold_paragraph(doc, "Token 重叠数", "两个文件中相同的 token（代码片段）数量。")
    add_bold_paragraph(doc, "开源占比", "相似行数在待检测文件总行数中的占比（抄袭比例）。")
    add_bold_paragraph(doc, "待检测文件总行数", "代码文件总行数。")
    add_bold_paragraph(doc, "相似行数", "被检测为相似或可疑的代码行数。")

    # 二、文件级代码检测结果
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("3. 文件级代码检测结果")
    set_run_font(h1_run, bold=True)

    for i, item in enumerate(stored_code_values, start=1):
        heading = doc.add_heading(f"第 {i} 条记录", level=3)
        for run in heading.runs:
            set_run_font(run, bold=True)

        add_field_value(doc, "待检测文件", item["code2"])
        add_field_value(doc, "参考文件", item["code3"])
        add_field_value(doc, "待检测文件的相似度得分", f"{item['code0'] * 100:.2f}%")
        add_field_value(doc, "参考文件的相似度得分", f"{item['code1'] * 100:.2f}%")
        add_field_value(doc, "Token 重叠数", item["code6"])
        add_field_value(doc, "相似行数", item["code9"])
        add_field_value(doc, "待检测文件总行数", item["code8"])
        add_field_value(doc, "开源占比", f"{item['code7'] * 100:.2f}%")

        # 添加解析段落
        analysis = doc.add_paragraph()
        analysis_run = analysis.add_run(
            f"解析：该比对显示待检测文件中有 {item['code0'] * 100:.2f}% 的内容与参考文件高度相似，"
            f"且参考文件中 {item['code1'] * 100:.2f}% 的代码出现在此文件中，应引起关注。"
        )
        set_run_font(analysis_run)

    # 三、字符级详细代码比对结果
    doc.add_paragraph()
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("4. 字符级详细代码比对结果")
    set_run_font(h1_run, bold=True)
    valid_matches = []

    # 改进数据处理逻辑，正确处理来自Java的ArrayList和String类型数据
    print(f"处理字符级详细代码比对数据，数据类型: {type(data)}")

    try:
        # 处理Java ArrayList类型数据
        if str(type(data)).find('java.util.ArrayList') >= 0:
            # 将每个Java字符串转换为Python字典
            processed_data = []
            for i in range(len(data)):
                item_str = str(data[i])
                try:
                    # 去除可能的前后引号
                    if item_str.startswith('"') and item_str.endswith('"'):
                        item_str = item_str[1:-1]
                    # 处理Java字符串中的转义字符
                    item_str = item_str.replace('\\"', '"')
                    item_dict = json.loads(item_str)
                    processed_data.append(item_dict)
                except json.JSONDecodeError as e:
                    print(f"无法解析Java字符串为JSON: {e}, 内容: {item_str[:100]}...")
            data = processed_data
            print(f"成功处理Java ArrayList数据，转换为{len(data)}个Python对象")

        # 确保data是一个列表
        if not isinstance(data, list):
            if isinstance(data, str):
                try:
                    data = json.loads(data)
                except:
                    data = [data]
            else:
                data = [data]

        # 遍历处理每个条目
        for entry in data:
            # 确保entry是字典类型
            if isinstance(entry, str):
                try:
                    entry = json.loads(entry)
                except Exception as e:
                    print(f"无法解析条目为JSON: {e}, 内容: {entry[:100]}...")
                    continue

            if not isinstance(entry, dict):
                print(f"跳过非字典类型的条目: {type(entry)}")
                continue

            for key, val in entry.items():
                if key == "sus_length" or key == "src_length":
                    continue
                if not val or not isinstance(val, dict) or "feature1" not in val:
                    continue

                if "suspicious-document" in key and "source-document" in key:
                    # 从键名中提取文件名
                    parts = key.split("source-document")
                    if len(parts) == 2:
                        sus_file = parts[0]
                        src_file = parts[1]
                    else:
                        sus_file = key.split("source-document")[0]
                        src_file = key.split("source-document")[1]
                else:
                    sus_file = "未知"
                    src_file = "未知"
                # 提取待检测文件名中的扩展名前 8 到 10 个字符，并保留扩展名
                match = re.search(r"([^\\\/]+)(\.[^\\\/]+)$", sus_file)  # 提取文件名和扩展名
                if match:
                    filename = match.group(1)  # 获取文件名部分（不含扩展名）
                    extension = match.group(2)  # 获取扩展名
                    sus_file = filename[-10:] if len(filename) > 10 else filename  # 截取最后 8 到 10 个字符
                    sus_file = "*" + sus_file + extension  # 添加扩展名

                feature = val["feature1"]
                sus_len = entry.get("sus_length", "未知")
                src_len = entry.get("src_length", "未知")

                valid_matches.append([
                    sus_file,
                    src_file,
                    feature["offset"],
                    feature["length"],
                    feature["srcOffset"],
                    feature["srcLength"],
                    sus_len,
                    src_len
                ])
    except Exception as e:
        print(f"处理字符级比对数据时出现严重错误: {e}")
        import traceback
        traceback.print_exc()

    if valid_matches:
        print(f"成功提取{len(valid_matches)}条字符级比对结果记录")
        
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        table.autofit = True  # ✅ 自动适应内容

        # 控制总宽度不超过 15cm
        col_widths = [Cm(2.5), Cm(2.5), Cm(1.5), Cm(1.5),
                    Cm(1.5), Cm(1.5), Cm(2), Cm(2)]

        headers = [
            "待检测文件", "参考文件",
            "可疑文件起始偏移", "可疑文件相似代码字符长度",
            "参考文件起始偏移", "参考文件相似代码字符长度",
            "待检测文件总字符数", "参考文件总字符数"
        ]

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header

            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "微软雅黑"
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))
            tcPr.append(parse_xml(r'<w:shd {} w:fill="4874CB"/>'.format(nsdecls("w"))))

        # 数据行
        for idx, row_data in enumerate(valid_matches):
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                cell = row_cells[i]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(value))
                set_run_font(run, size=Pt(9))  # ✅ 小一点字体

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))

                # ✅ 移除 noWrap 的处理，不添加即可自动换行！

                # 设置背景色
                shade_color = "CFD6EC" if idx % 2 == 0 else "E9ECF6"
                tcPr.append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{shade_color}"/>'))

    else:
        para = doc.add_paragraph()
        run = para.add_run("未检测到任何字符级相似代码片段。")
        set_run_font(run, size=Pt(12))


    # 保存报告
    word_output_file = output_file
    pdf_output_file = output_file.replace('.docx', '.pdf')
    if output_file.endswith('.pdf'):
        word_output_file = output_file.replace('.pdf', '.docx')
        pdf_output_file = output_file
    
    try:
        # 确保目录存在
        output_dir = os.path.dirname(word_output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 先保存为 Word 文档
        doc.save(word_output_file)
        print(f"Word报告成功保存到: {word_output_file}")
        
        # 尝试转换为 PDF
        pdf_converted = convert_word_to_pdf(word_output_file, pdf_output_file)
        
        if pdf_converted:
            print(f"PDF报告成功保存到: {pdf_output_file}")
            return pdf_output_file if output_file.endswith('.pdf') else word_output_file
        else:
            print(f"警告: 无法转换为PDF格式，请确保安装了Microsoft Word或其他兼容工具")
            return word_output_file
    except Exception as e:
        print(f"保存文档时出错: {e}")
        import traceback
        traceback.print_exc()
    
    return word_output_file


# 示例调用
if __name__ == "__main__":
    # 测试是否可以生成PDF
    try:
        test_docx = os.path.join(current_dir, "test_doc.docx")
        test_pdf = os.path.join(current_dir, "test_pdf.pdf")
        
        # 创建一个简单的测试文档
        doc = Document()
        doc.add_paragraph("这是一个测试文档")
        doc.save(test_docx)
        
        # 测试转换
        result = convert_word_to_pdf(test_docx, test_pdf)
        print(f"PDF转换测试结果: {result}")
        
        # 清理测试文件
        if os.path.exists(test_docx):
            os.remove(test_docx)
        if os.path.exists(test_pdf):
            os.remove(test_pdf)
    except Exception as e:
        print(f"PDF转换测试失败: {e}")
        
    print(f"如要使用完整功能，请在命令行运行: pip install docx2pdf pywin32 comtypes")
