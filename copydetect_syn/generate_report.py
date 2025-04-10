from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os,re
from datetime import datetime
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
from config import WORD_REPORT_NAME # 配置文件

current_dir = os.path.dirname(os.path.abspath(__file__))

def create_government_blockchain_report(
    stored_code_values,
    data,
    output_file=os.path.join(current_dir,'document',WORD_REPORT_NAME)
):
    # 若 data 是字符串则先解析
    if isinstance(data, str):
        try:
            data = json.loads(data)

            # 如果解析后是 list[str]，还要进一步把每个 entry 再做 json.loads
            if isinstance(data, list) and all(isinstance(d, str) for d in data):
                data = [json.loads(d) for d in data]

        except Exception as e:
            print(f"解析 data 出错，无法解析为 JSON：{e}")
            return

    # 创建新文档
    doc = Document()

    # 设置正文中文字体为宋体，英文字体为Times New Roman
    def set_run_font(run, size=Pt(12), bold=False):
        font = run.font
        font.size = size
        font.bold = bold

        # 设置英文字体
        run.font.name = "Times New Roman"

        # 设置中文字体为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # === 第一页：封面 ===
    # 添加页眉标题
    header_para = doc.sections[0].header.paragraphs[0]
    header_run = header_para.add_run("源代码自研率检测报告")
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header_run, size=Pt(12))

    # 添加主标题（居中）
    for _ in range(8):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = main_title.add_run("源代码自研率检测报告")
    set_run_font(title_run, size=Pt(28), bold=True)
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑") # 这个标题单独设置成微软雅黑吧

    # 底部信息
    for _ in range(12):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    bottom_info = doc.add_paragraph()
    bottom_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_time = datetime.now().strftime("%Y 年 %m 月 %d 日 %H:%M:%S")

    bottom_run_first = bottom_info.add_run("编制单位：佛山大学\n")
    set_run_font(bottom_run_first, size=Pt(14))
    bottom_run_second = bottom_info.add_run(f"编制时间：{current_time}")
    set_run_font(bottom_run_second, size=Pt(14))

    # 添加分页符
    doc.add_page_break()

    # === 第二页：内容 ===
    # 设置标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("源代码自研率检测报告")
    set_run_font(title_run, size=Pt(16), bold=True)

    # 添加简介文本
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "本报告对上传代码自研率进行了检测，包括代码同源分析、代码引用、代码行级开源占比和代码文件级开源占比。"
    )
    set_run_font(intro_run)

    # 添加定义标题
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. 定义")
    set_run_font(h1_run, bold=True)

    # 添加项目简介内容
    intro_text = """源代码安全审计的主要目的是提高源代码质量，通过对程序源代码进行检查和分析，发现源代码在软件设计、测试、应用部署等各阶段中可能存在的安全缺陷或安全漏洞，从源头上避免潜在的安全风险。"""
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(intro_text)
    set_run_font(intro_run)

    # 一、字段说明
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("2、字段说明")
    set_run_font(h1_run, bold=True)

    doc.add_paragraph("以下是检测系统输出的字段说明：")

    def add_bold_paragraph(doc, field_title, description):
        p = doc.add_paragraph()
        run1 = p.add_run(f"{field_title}：")
        run1.bold = True
        run2 = p.add_run(description)

    add_bold_paragraph(doc, "待检测代码路径", "用户上传、参与查重的代码文件路径。")
    add_bold_paragraph(doc, "参考代码路径", "系统比对的原始或参考源代码路径。")
    add_bold_paragraph(doc, "待检测文件的相似度得分", "该文件中与参考代码相似的比例（0~1）。")
    add_bold_paragraph(doc, "参考文件的相似度得分", "参考代码中出现在待检测文件中的比例（0~1）。")
    add_bold_paragraph(doc, "Token 重叠数", "两个文件中相同的 token（代码片段）数量。")
    add_bold_paragraph(doc, "开源占比", "相似行数在待检测文件总行数中的占比（抄袭比例）。")
    add_bold_paragraph(doc, "待检测文件总行数", "代码文件总行数。")
    add_bold_paragraph(doc, "相似行数", "被检测为相似或可疑的代码行数。")

    # 二、文件级代码检测结果
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("3、文件级代码检测结果")
    set_run_font(h1_run, bold=True)

    for i, item in enumerate(stored_code_values, start=1):
        doc.add_heading(f"第 {i} 条记录", level=3)
        doc.add_paragraph(f"待检测文件：{item['code2']}")
        doc.add_paragraph(f"参考文件：{item['code3']}")
        doc.add_paragraph(f"待检测文件的相似度得分：{item['code0'] * 100:.2f}%")
        doc.add_paragraph(f"参考文件的相似度得分：{item['code1'] * 100:.2f}%")
        doc.add_paragraph(f"Token 重叠数：{item['code6']}")
        doc.add_paragraph(f"相似行数：{item['code9']}")
        doc.add_paragraph(f"待检测文件总行数：{item['code8']}")
        doc.add_paragraph(f"开源占比：{item['code7'] * 100:.2f}%")

        doc.add_paragraph(
            f"解析：该比对显示待检测文件中有 {item['code0'] * 100:.2f}% 的内容与参考文件高度相似，"
            f"且参考文件中 {item['code1'] * 100:.2f}% 的代码出现在此文件中，应引起关注。"
        )

    # 三、字符级详细代码比对结果
    doc.add_paragraph()
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("4、字符级详细代码比对结果")
    set_run_font(h1_run, bold=True)
    valid_matches = []

    # 改进数据处理逻辑，正确处理来自Java的ArrayList和String类型数据
    print(f"处理字符级详细代码比对数据，数据类型: {type(data)}")

    try:
        # 处理Java ArrayList类型数据
        if str(type(data)).find('java.util.ArrayList') >= 0:
            # 将每个Java字符串转换为Python字典
            processed_data = []
            for i in range(len(data)):
                item_str = str(data[i])
                try:
                    # 去除可能的前后引号
                    if item_str.startswith('"') and item_str.endswith('"'):
                        item_str = item_str[1:-1]
                    # 处理Java字符串中的转义字符
                    item_str = item_str.replace('\\"', '"')
                    item_dict = json.loads(item_str)
                    processed_data.append(item_dict)
                except json.JSONDecodeError as e:
                    print(f"无法解析Java字符串为JSON: {e}, 内容: {item_str[:100]}...")
            data = processed_data
            print(f"成功处理Java ArrayList数据，转换为{len(data)}个Python对象")

        # 确保data是一个列表
        if not isinstance(data, list):
            if isinstance(data, str):
                try:
                    data = json.loads(data)
                except:
                    data = [data]
            else:
                data = [data]

        # 遍历处理每个条目
        for entry in data:
            # 确保entry是字典类型
            if isinstance(entry, str):
                try:
                    entry = json.loads(entry)
                except Exception as e:
                    print(f"无法解析条目为JSON: {e}, 内容: {entry[:100]}...")
                    continue

            if not isinstance(entry, dict):
                print(f"跳过非字典类型的条目: {type(entry)}")
                continue

            for key, val in entry.items():
                if key == "sus_length" or key == "src_length":
                    continue
                if not val or not isinstance(val, dict) or "feature1" not in val:
                    continue

                if "suspicious-document" in key and "source-document" in key:
                    # 从键名中提取文件名
                    parts = key.split("source-document")
                    if len(parts) == 2:
                        sus_file = parts[0]
                        src_file = parts[1]
                    else:
                        sus_file = key.split("source-document")[0]
                        src_file = key.split("source-document")[1]
                else:
                    sus_file = "未知"
                    src_file = "未知"
                # 提取待检测文件名中的扩展名前 8 到 10 个字符，并保留扩展名
                match = re.search(r"([^\\\/]+)(\.[^\\\/]+)$", sus_file)  # 提取文件名和扩展名
                if match:
                    filename = match.group(1)  # 获取文件名部分（不含扩展名）
                    extension = match.group(2)  # 获取扩展名
                    sus_file = filename[-10:] if len(filename) > 10 else filename  # 截取最后 8 到 10 个字符
                    sus_file = "*" + sus_file + extension  # 添加扩展名

                feature = val["feature1"]
                sus_len = entry.get("sus_length", "未知")
                src_len = entry.get("src_length", "未知")

                valid_matches.append([
                    sus_file,
                    src_file,
                    feature["offset"],
                    feature["length"],
                    feature["srcOffset"],
                    feature["srcLength"],
                    sus_len,
                    src_len
                ])
    except Exception as e:
        print(f"处理字符级比对数据时出现严重错误: {e}")
        import traceback
        traceback.print_exc()

    if valid_matches:
        print(f"成功提取{len(valid_matches)}条字符级比对结果记录")
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        table.autofit = False

        col_widths = [Inches(2.2), Inches(2.2), Inches(1), Inches(1),
                    Inches(1), Inches(1), Inches(1), Inches(1)]

        headers = [
            "待检测文件", "参考文件",
            "可疑文件起始偏移", "可疑文件相似代码字符长度",
            "参考文件起始偏移", "参考文件相似代码字符长度",
            "待检测文件总字符数", "参考文件总字符数"
        ]

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header
            cell.width = col_widths[i]

            # 设置字体样式
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # 白色字体
            # 表头中文字体设置为微软雅黑
            run.font.name = "微软雅黑"
            run.font.size = Pt(9)  # 设置字体大小为9pt
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.bold = True  # 设置加粗

            # 设置水平居中对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 设置垂直居中对齐
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vertical_alignment = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
            tcPr.append(vertical_alignment)

            # 设置表头背景色（PowerPoint默认的蓝色）
            cell._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="4874CB"/>'.format(nsdecls("w")))
            )

        # 添加数据行
        for idx, row_data in enumerate(valid_matches):
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                cell = row_cells[i]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(value))  # 添加单元格内容
                set_run_font(run, size=Pt(11))  # 使用 set_run_font 设置字体
                cell.width = col_widths[i]

                # 设置水平居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 设置垂直居中对齐
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # 创建垂直居中XML元素
                vertical_alignment = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
                tcPr.append(vertical_alignment)

                # 允许自动换行
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(OxmlElement("w:noWrap"))
                tcPr.remove(tcPr.xpath("w:noWrap")[0])  # 删除不换行限制

                # 背景色用powerpoint的默认颜色
                if idx % 2 == 0:  # 奇数数据行背景色
                    cell._tc.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="CFD6EC"/>'.format(nsdecls("w")))
                    )
                else:  # 偶数行背景色
                    cell._tc.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="E9ECF6"/>'.format(nsdecls("w")))
                    )
    else:
        doc.add_paragraph("未检测到任何字符级相似代码片段。")

    # 保存报告
    try:
        output_dir = os.path.dirname(output_file)
        if (output_dir and not os.path.exists(output_dir)):
            os.makedirs(output_dir)
        doc.save(output_file)
        print(f"报告成功保存到: {output_file}")
    except Exception as e:
        print(f"保存文档时出错: {e}")
    return output_file


# 示例调用
if __name__ == "__main__":
    # stored_code_values = [
    #     {
    #         'code0': 0.4228965112870126,
    #         'code1': 0.9918426103646834,
    #         'code2': 'e:\\EL\\Server_Project\\Server_Project_hub\\copydetect_syn\\uploads\\suspicious-document000001.java',
    #         'code3': 'e:\\EL\\Server_Project\\Server_Project_hub\\copydetect_syn\\downs\\source-document0172528.java',
    #         'code6': 2885,
    #         'code7': 0.5289256198347108,
    #         'code8': 242,
    #         'code9': 128
    #     },
    #     {
    #         'code0': 0.4183523893286426,
    #         'code1': 0.9995872884853487,
    #         'code2': 'e:\\EL\\Server_Project\\Server_Project_hub\\copydetect_syn\\uploads\\suspicious-document000001.java',
    #         'code3': 'e:\\EL\\Server_Project\\Server_Project_hub\\copydetect_syn\\downs\\source-document1070483.java',
    #         'code6': 2854,
    #         'code7': 0.30165289256198347,
    #         'code8': 242,
    #         'code9': 73
    #     }
    # ]
    
    # data  =   [{"suspicious-document000001.javasource-document1070483.java":
    #             {"feature1":{"offset":12,"length":2854,"srcOffset":0,"srcLength":2422}},"sus_length":6822,"src_length":2423}, 
    #         {"suspicious-document000001.javasource-document2264813.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":0,"srcLength":2159}},"sus_length":6822,"src_length":2228}, 
    #         {"suspicious-document000001.javasource-document2264819.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":0,"srcLength":1871}},"sus_length":6822,"src_length":1940},
    #         {"suspicious-document000001.javasource-document0172528.java":
    #             {"feature1":{"offset":1688,"length":2885,"srcOffset":0,"srcLength":2067}},"sus_length":6822,"src_length":2084}, 
    #         {"suspicious-document000001.javasource-document2264812.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":295,"srcLength":2302}},"sus_length":6822,"src_length":2696}, 
    #         {"suspicious-document000001.javasource-document2264820.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":0,"srcLength":1767}},"sus_length":6822,"src_length":1866}, 
    #         {"suspicious-document000001.javasource-document2264815.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":380,"srcLength":5235}},"sus_length":6822,"src_length":5684}, 
    #         {"suspicious-document000001.javasource-document2264818.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":384,"srcLength":6339}},"sus_length":6822,"src_length":6794}, 
    #         {"suspicious-document000001.javasource-document2264814.java":
    #             {"feature1":{"offset":4573,"length":2178,"srcOffset":1347,"srcLength":3718}},"sus_length":6822,"src_length":6823}, 
    #         {"suspicious-document000001.javasource-document0172522.java":
    #             {},"sus_length":6822,"src_length":335}, 
    #         {"suspicious-document000001.javasource-document0172544.java":
    #             {},"sus_length":6822,"src_length":393}
    #         ]
    
    
    # output_file = create_government_blockchain_report(stored_code_values, data=data)
    print(f"报告已生成: ")
