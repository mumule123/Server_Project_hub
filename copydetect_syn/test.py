from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn


def create_government_blockchain_report(
    output_file="政务区块链基础平台项目源代码自研率检测报告.docx",
):
    # 创建新文档
    doc = Document()

    # 设置正文中文字体为宋体，英文字体为Times New Roman
    def set_run_font(run, size=Pt(12), bold=False):
        font = run.font
        font.size = size
        font.bold = bold

        # 设置英文字体
        run.font.name = "Times New Roman"

        # 设置中文字体为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # === 第一页：封面 ===
    # 添加页眉标题
    header_para = doc.sections[0].header.paragraphs[0]
    header_run = header_para.add_run("《政务区块链基础平台》项目源代码自研率检测报告")
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header_run, size=Pt(12))

    # 添加主标题（居中）
    for _ in range(8):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = main_title.add_run("《政务区块链基础平台》项目\n源代码自研率检测报告")
    set_run_font(title_run, size=Pt(22), bold=True)

    # 添加底部信息
    for _ in range(12):  # 添加空行
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    bottom_info = doc.add_paragraph()
    bottom_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_run = bottom_info.add_run("编制单位：佛山大学\n2024 年 10 月")
    set_run_font(bottom_run, size=Pt(14))

    # 添加分页符
    doc.add_page_break()

    # === 第二页：内容 ===
    # 设置标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("《政务区块链基础平台》项目源代码自研率检测报告")
    set_run_font(title_run, size=Pt(16), bold=True)

    # 添加简介文本
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "本报告对政务区块链基础平台项目代码自研率进行了检测，包括代码同源分析、代码引用、代码行级开源占比和代码文件级开源占比。"
    )
    set_run_font(intro_run)

    # 添加定义标题
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. 定义")
    set_run_font(h1_run, bold=True)


    # 添加项目简介内容
    intro_text = """政务区块链基础平台借助区块链等前沿信息技术，发挥其不可更改和追溯的特性，构建一个集成的平台接口/服务系统，实现数据和文件的链上存储，确保数据的可信查询和安全使用。同时，积极探索区块链数据交换和多部门间的业务合作，以创新数据管理的新方法，提升电子政务的服务和管理能力，提高政务工作的效率，并促进政府数据的有效治理。"""
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(intro_text)
    set_run_font(intro_run)

    # 添加表格标题
    table_title = doc.add_paragraph()
    table_title_run = table_title.add_run("表 1 项目模块结构")
    set_run_font(table_title_run)

    # 创建表格
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # 设置表头
    header_cells = table.rows[0].cells
    for i, text in enumerate(["序号", "子系统名称", "子模块名称"]):
        run = header_cells[i].paragraphs[0].add_run(text)
        set_run_font(run, bold=True)

    # 添加数据
    modules = [
        (
            "1",
            "EasySpiral",
            [
                ("区块链管理", ["节点管理", "通道管理", "应用管理", "合约管理"]),
                (
                    "区块链应用/数字证书",
                    ["证书颁发", "证书管理", "证书使用", "证书颁发查询"],
                ),
                (
                    "区块链应用/电子印章",
                    [
                        "印章查询",
                        "印章服务",
                        "盖章记录管理",
                        "盖章规则配置",
                        "印章权限管理",
                        "印章变更记录",
                        "印章验真",
                    ],
                ),
                ("区块链应用/存", ["存证操作"]),
            ],
        )
    ]

    # 填充表格数据
    for module in modules:
        for submodule_group in module[2]:
            main_row = table.add_row()
            if module[2].index(submodule_group) == 0:
                run = main_row.cells[0].paragraphs[0].add_run(module[0])
                set_run_font(run)
                run = main_row.cells[1].paragraphs[0].add_run(module[1])
                set_run_font(run)
            run = main_row.cells[2].paragraphs[0].add_run(submodule_group[0])
            set_run_font(run)

            for submodule in submodule_group[1]:
                sub_row = table.add_row()
                run = sub_row.cells[2].paragraphs[0].add_run(submodule)
                set_run_font(run)

    # 保存文档
    doc.save(output_file)
    return output_file


if __name__ == "__main__":
    output_file = create_government_blockchain_report()
    print(f"报告已生成: {output_file}")